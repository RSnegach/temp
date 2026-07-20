# -*- coding: utf-8 -*-
"""
Build the hardened contract as full prose: a Master Services Agreement plus
Amendment No. 1 plus Exhibit B. The obligations are embedded in ordinary
contract language, not a clause list, so the analyst must FIND them. Traps:
  - Supersession stated obliquely: "Section 7.2 is deleted in its entirety and
    replaced by the following" (the 60-business-day notice becomes 90).
  - A chained deadline: the integration report is due "within twenty (20)
    business days after the Client's sign-off of the Statement of Work" (which
    is itself Section 4.1's 30-calendar-day obligation).
  - A second conditional whose FACT lives in Exhibit B: the security review
    applies "where the data classification in Exhibit B is High"; Exhibit B, far
    below, lists the classification as High.
  - A conditional that stays OFF: the breach drill applies only if the Provider
    "Processes Personal Data"; a recital states it does not.

All obligations and facts match contract.py so the golden ties out. Black text,
neutral headers, no em dashes.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0, 0, 0)

def newdoc():
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(11); n.font.color.rgb = BLACK
    n.paragraph_format.space_after = Pt(6); n.paragraph_format.line_spacing = 1.15
    rpr = n.element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts(); rf.set(qn('w:eastAsia'), 'Calibri')
    for s in doc.sections:
        s.top_margin = Inches(1.0); s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)
    return doc

def sr(r, size=11, bold=False, italic=False):
    r.font.name = "Calibri"; r.font.size = Pt(size); r.font.bold = bold
    r.font.italic = italic; r.font.color.rgb = BLACK

def para(doc, t, size=11, bold=False, after=6, align=None, italic=False):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    if align: p.alignment = align
    r = p.add_run(t); sr(r, size, bold, italic); return p

def h(doc, t):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(t); sr(r, 12, True); return p

def main(out="MSA_and_Amendment.docx"):
    doc = newdoc()
    para(doc, "MASTER SERVICES AGREEMENT", 16, True, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Between Northwind Analytics Inc. (Client) and Meridian Technical Services LLC (Provider)",
         10.5, False, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Contract No. MER-2025-0087", 10, False, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    h(doc, "Recitals")
    para(doc, "This Master Services Agreement (the \"Agreement\") is entered into and effective as of "
              "January 15, 2025 (the \"Effective Date\"). The Provider will supply technical and analytical "
              "services as described in one or more Statements of Work. The Provider does not Process "
              "Personal Data on behalf of the Client under this Agreement, and the parties acknowledge that "
              "the data-protection operational obligations of Section 9 apply only if that circumstance "
              "changes. Capitalized terms not defined inline have the meanings given in Exhibit B.")

    h(doc, "Section 1. Mobilization")
    para(doc, "1.1 The Provider shall convene a project kickoff meeting with the Client within ten (10) "
              "calendar days after the Effective Date. 1.2 The Provider shall deliver its first monthly "
              "invoice within five (5) business days after the Effective Date. 1.3 The Provider shall deliver "
              "a certificate of insurance evidencing the required coverage within fifteen (15) calendar days "
              "after the Effective Date.")

    h(doc, "Section 2. Reporting")
    para(doc, "2.1 Beginning in the month after the Effective Date, the Provider shall deliver a written "
              "status report by the fifth (5th) day of each month through the end of the calendar year. "
              "2.2 The Provider and the Client shall hold a quarterly business review no later than three (3) "
              "business days before the end of each calendar quarter.")

    h(doc, "Section 3. Fees and Audit")
    para(doc, "3.1 The Client shall pay the fees set out in the applicable Statement of Work. 3.2 If the "
              "total annual fees payable under this Agreement exceed two hundred fifty thousand dollars "
              "($250,000), the Provider shall deliver an annual audit support package no later than July 31 "
              "of the contract year. The current fee schedule in Exhibit B reflects the annualized amount.")

    h(doc, "Section 4. Statement of Work and Deliverables")
    para(doc, "4.1 The Client shall sign off on the initial Statement of Work within thirty (30) calendar "
              "days after the Effective Date. 4.2 The Provider shall deliver the integration report within "
              "twenty (20) business days after the Client's sign-off of the initial Statement of Work under "
              "Section 4.1. For the avoidance of doubt, the integration report deadline is measured from the "
              "date the sign-off obligation in Section 4.1 comes due.")

    h(doc, "Section 7. Term, Renewal, and Notice")
    para(doc, "7.1 This Agreement has an initial term through December 31 of the contract year and renews "
              "automatically unless a party gives timely notice of non-renewal. 7.2 The last day on which a "
              "party may give written notice of non-renewal is the date sixty (60) business days before "
              "December 31; a party intending not to renew must give that notice on or before that deadline, "
              "and the parties shall calendar that deadline date whether or not either party has yet decided "
              "to renew.")

    h(doc, "Section 9. Data Protection")
    para(doc, "9.1 If and only if the Provider Processes Personal Data on behalf of the Client, the Provider "
              "shall conduct a data-protection breach response drill within sixty (60) calendar days after "
              "the Effective Date. 9.2 Where the data classification stated in Exhibit B is \"High\", the "
              "Provider shall complete an independent security review within seventy-five (75) calendar days "
              "after the Effective Date, regardless of whether it Processes Personal Data.")

    h(doc, "Section 12. Rolling of Dates")
    para(doc, "12.1 If any deadline stated in calendar days, or any fixed calendar date, falls on a Saturday, "
              "Sunday, or a holiday listed in Exhibit B, the deadline is the next business day. 12.2 A period "
              "stated in business days is counted using only business days and is not further adjusted.")

    # ---------------- Amendment ----------------
    doc.add_page_break()
    para(doc, "AMENDMENT NO. 1 TO THE MASTER SERVICES AGREEMENT", 14, True, after=2,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Effective March 3, 2025", 10.5, False, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "The parties amend the Agreement as follows. Except as amended below, the Agreement remains "
              "in full force.")
    para(doc, "A. Section 7.2 is deleted in its entirety and replaced by the following: \"The last day on "
              "which a party may give written notice of non-renewal is the date ninety (90) business days "
              "before December 31; the parties shall calendar that deadline date whether or not either party "
              "has yet decided to renew.\"")
    para(doc, "B. A new Section 4.3 is added: \"The Provider shall deliver a termination true-up statement "
              "within forty-five (45) calendar days after the effective date of Amendment No. 1.\"")

    # ---------------- Exhibit B ----------------
    doc.add_page_break()
    para(doc, "EXHIBIT B  Definitions, Fee Schedule, Data Classification, and Holidays", 13, True, after=8)
    para(doc, "Definitions. \"Process Personal Data\" has the meaning given under applicable data-protection "
              "law. \"Business day\" means a weekday that is not a holiday listed below.")
    para(doc, "Fee schedule. The annualized total fees payable under Contract No. MER-2025-0087 are "
              "$312,000.")
    para(doc, "Data classification. The data handled under this engagement is classified as: High.")
    para(doc, "Holiday calendar (observed, contract year). New Year's Day Jan 1; Martin Luther King Jr. Day "
              "Jan 20; Presidents Day Feb 17; Memorial Day May 26; Juneteenth Jun 19; Independence Day Jul 4; "
              "Labor Day Sep 1; Thanksgiving Nov 27 and the day after Nov 28; Christmas Day Dec 25.")

    doc.save(out)
    print("wrote", out)

if __name__ == "__main__":
    main()
