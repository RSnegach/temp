# -*- coding: utf-8 -*-
"""Build the three DFMEA input files. Distributed: spec (function context),
ratings+AP doc (how to score), failure-mode xlsx (S,O given + controls + dup/superseded)."""
import json
import dfmea as D
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

BLACK=RGBColor(0,0,0); ACC=RGBColor(0x2A,0x4A,0x6B); GREY=RGBColor(0x44,0x44,0x44)

# ---------- docx helpers ----------
def newdoc():
    doc=Document(); n=doc.styles["Normal"]; n.font.name="Calibri"; n.font.size=Pt(11); n.font.color.rgb=BLACK
    n.paragraph_format.space_after=Pt(6); n.paragraph_format.line_spacing=1.15
    rpr=n.element.get_or_add_rPr(); rf=rpr.get_or_add_rFonts(); rf.set(qn('w:eastAsia'),'Calibri')
    for s in doc.sections:
        s.top_margin=Inches(0.9); s.bottom_margin=Inches(0.9); s.left_margin=Inches(1.0); s.right_margin=Inches(1.0)
    return doc
def sr(r,size=11,bold=False,color=BLACK,italic=False):
    r.font.name="Calibri"; r.font.size=Pt(size); r.font.bold=bold; r.font.italic=italic; r.font.color.rgb=color
    rpr=r._element.get_or_add_rPr(); rf=rpr.get_or_add_rFonts(); rf.set(qn('w:eastAsia'),'Calibri')
def para(doc,t,size=11,bold=False,color=BLACK,after=6):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after); r=p.add_run(t); sr(r,size,bold,color); return p
def heading(doc,num,text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(f"{num}   {text}"); sr(r,13,True,ACC)
    pPr=p._p.get_or_add_pPr(); pb=OxmlElement('w:pBdr'); b=OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6'); b.set(qn('w:space'),'2'); b.set(qn('w:color'),'2A4A6B')
    pb.append(b); pPr.append(pb); return p
def bullet(doc,t,size=11):
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3); r=p.add_run(t); sr(r,size); return p
def setbg(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hexc); tcPr.append(shd)
def table(doc,headers,rows,widths,fs=9):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; r=c.paragraphs[0].add_run(h); sr(r,fs,True,RGBColor(0xFF,0xFF,0xFF)); setbg(c,"2A4A6B")
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for ci,val in enumerate(row):
            cells[ci].text=""; r=cells[ci].paragraphs[0].add_run(str(val)); sr(r,fs)
            if ri%2==1: setbg(cells[ci],"EEF2F6")
    for row in t.rows:
        for ci,w in enumerate(widths): row.cells[ci].width=Inches(w)
    return t

# =========================================================
# FILE 1: Fuel Delivery System DFMEA Scope and Function
# =========================================================
doc=newdoc()
p=doc.add_paragraph(); r=p.add_run("PRODUCT ENGINEERING - RELIABILITY"); sr(r,9,True,GREY)
para(doc,"Fuel Delivery System - Design FMEA Scope",17,True,ACC,after=2)
para(doc,"Program FX-mid gasoline PFI  |  Subsystem definition and functions",11,False,GREY,after=2)
para(doc,"DOC FDS-DFMEA-SCOPE Rev A",9.5,False,GREY,after=10)
para(doc,"This document defines the fuel delivery subsystem for the Design FMEA: its items, "
     "their functions, and the intended performance. It does not assign risk ratings or list "
     "failure modes. Ratings method and thresholds are in the FMEA Ratings and Action Priority "
     "Reference (FDS-DFMEA-RATE). The failure-mode working list with severity, occurrence, and "
     "current detection controls is in the failure-mode register (FDS-DFMEA-MODES.xlsx). The DFMEA "
     "is completed by scoring that register against the reference.")
heading(doc,"1","Subsystem Boundary")
para(doc,"The fuel delivery subsystem carries gasoline from the tank to the injectors at regulated "
     "pressure and meters it into the intake port on command. In scope: fuel pump, filter, lines and "
     "quick connectors, rail, pressure regulator and its diaphragm, injectors and their seals, and tank "
     "sealing at the pump flange. Out of scope: tank structure, evap canister, engine control strategy.")
heading(doc,"2","Items and Functions")
table(doc,["Item","Primary function","Performance requirement"],
 [["Fuel Pump","Deliver rated flow and pressure","Maintain rail pressure across the flow range"],
  ["Fuel Filter","Filter particulate","Protect injectors, hold flow across service life"],
  ["Fuel Rail","Contain pressurized fuel, distribute to injectors","No external leak at rated cycling"],
  ["Quick Connector","Join fuel lines","Retain and seal under vibration"],
  ["Pressure Regulator","Regulate rail pressure","Hold set pressure across load"],
  ["Regulator Diaphragm","Seal reference chamber","No rupture across durability"],
  ["Injector","Meter fuel per command","Deliver commanded quantity, seal at seat"],
  ["Tank Seal","Contain fuel vapor and liquid at pump flange","No weep or evap leak"]],
 [1.7,2.6,2.7])
heading(doc,"3","Failure Consequences of Interest")
para(doc,"Failure effects range from driveability and emissions concerns to safety. Any external fuel "
     "leak or unseated connection is treated as a fire-risk safety effect and carries the highest severity. "
     "The severity value for each failure mode is provided in the register; do not re-rate severity here.")
for section in doc.sections:
    fp=section.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=fp.add_run("FDS-DFMEA-SCOPE Rev A   |   Page "); sr(r,8,False,GREY)
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); fp._p.append(fld)
doc.core_properties.author="Product Engineering"; doc.core_properties.title="Fuel Delivery System DFMEA Scope"
doc.save("inputs/FDS-DFMEA-SCOPE_Subsystem_and_Functions.docx"); print("saved scope doc")

# =========================================================
# FILE 2: Ratings guide + AP table + action rule
# =========================================================
doc=newdoc()
p=doc.add_paragraph(); r=p.add_run("PRODUCT ENGINEERING - RELIABILITY"); sr(r,9,True,GREY)
para(doc,"FMEA Ratings and Action Priority Reference",17,True,ACC,after=2)
para(doc,"AIAG-VDA method (2019)  |  Detection rating map and Action Priority table",11,False,GREY,after=2)
para(doc,"DOC FDS-DFMEA-RATE Rev A",9.5,False,GREY,after=10)
para(doc,"This reference gives the rating method for the Design FMEA. Severity and Occurrence values "
     "are supplied per failure mode in the register. Detection is not supplied; derive it from the current "
     "detection control named in the register using the Detection map below. Then determine Action Priority "
     "from the Action Priority table. Action Priority replaces RPN; do not multiply S, O, and D.")
heading(doc,"1","Detection Rating from Current Control")
para(doc,"Look up the failure mode's current detection control in the register, then read its Detection "
     "rating here. Lower detection numbers mean better detection.")
drows=sorted(D.DETECTION_MAP.items(), key=lambda kv:kv[1])
table(doc,["Current detection control","Detection rating (D)"],
      [[k,str(v)] for k,v in drows],[4.6,1.9])
heading(doc,"2","Action Priority Table")
para(doc,"Action Priority is High, Medium, or Low, determined by the combination of Severity, "
     "Occurrence, and Detection. Find the Severity band, then the Occurrence band, then read the "
     "Detection column. This is a lookup, not a calculation.")
# build the compact AP table exactly as engine prints it
Sband=[(9,10),(7,8),(4,6),(2,3),(1,1)]; Oband=[(6,10),(4,5),(2,3),(1,1)]; Dband=[(1,4),(5,7),(8,10)]
aprows=[]
for slo,shi in Sband:
    for olo,ohi in Oband:
        cells=[D.action_priority(shi,ohi,dh) for (dl,dh) in Dband]
        srange=f"{slo}-{shi}" if slo!=shi else f"{slo}"
        orange=f"{olo}-{ohi}" if olo!=ohi else f"{olo}"
        aprows.append([srange,orange,cells[0],cells[1],cells[2]])
table(doc,["Severity","Occurrence","D 1-4","D 5-7","D 8-10"],aprows,[1.4,1.5,1.3,1.3,1.4],fs=9)
para(doc,"Severity 1 is Low for every Occurrence and Detection.",9.5,False,GREY)
heading(doc,"3","Action Rule")
bullet(doc,"Action Priority High: action is required. Enter Yes in the action column.")
bullet(doc,"Action Priority Medium: review and justify. Enter Review.")
bullet(doc,"Action Priority Low: acceptable as-is. Enter No.")
heading(doc,"4","Ranking")
para(doc,"Rank the completed failure modes by Action Priority (High, then Medium, then Low), and within "
     "each priority by descending Severity, then Occurrence, then Detection.")
for section in doc.sections:
    fp=section.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=fp.add_run("FDS-DFMEA-RATE Rev A   |   Page "); sr(r,8,False,GREY)
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); fp._p.append(fld)
doc.core_properties.author="Product Engineering"; doc.core_properties.title="FMEA Ratings and Action Priority Reference"
doc.save("inputs/FDS-DFMEA-RATE_Ratings_and_Action_Priority.docx"); print("saved ratings doc")

# =========================================================
# FILE 3: failure-mode register (xlsx) with dup/superseded
# =========================================================
NAVY_OK="000000"; WHITE="FFFFFF"; ACCENTX="2A4A6B"
thin=Side(style="thin",color="BBBBBB"); B=Border(left=thin,right=thin,top=thin,bottom=thin)
wb=Workbook(); ws=wb.active; ws.title="Failure_Modes"
ws["A1"]="Fuel Delivery System - Failure Mode Register (Rev C working list)"
ws["A1"].font=Font(name="Calibri",size=13,bold=True,color="2A4A6B")
hdrs=["fm_id","rev","status","item","function","failure_mode","effect","severity","occurrence","current_detection_control"]
for i,h in enumerate(hdrs):
    c=ws.cell(2,1+i,h); c.fill=PatternFill("solid",fgColor=NAVY_OK); c.font=Font(name="Calibri",size=10,bold=True,color=WHITE)
    c.alignment=Alignment(horizontal="center",vertical="center",wrap_text=True); c.border=B
r=3
for (fid,rev,status,item,func,fm,eff,S,O,control) in D.MODES:
    vals=[fid,rev,status,item,func,fm,eff,S,O,control]
    for ci,v in enumerate(vals):
        c=ws.cell(r,1+ci,v); c.font=Font(name="Calibri",size=9); c.border=B
        c.alignment=Alignment(horizontal="center" if ci in (0,1,2,7,8) else "left",vertical="center",wrap_text=True)
    r+=1
for col,w in zip("ABCDEFGHIJ",[7,5,13,15,22,26,26,9,11,26]): ws.column_dimensions[col].width=w
ws.freeze_panes="A3"
wb.properties.creator="Product Engineering"; wb.properties.title="Fuel Delivery Failure Mode Register"
REG="inputs/FDS-DFMEA-MODES_Failure_Mode_Register.xlsx"
wb.save(REG)
from xlsx_live import normalize_decimals, set_excel_fingerprint
normalize_decimals(REG); set_excel_fingerprint(REG, creator="Product Engineering")
print("saved failure-mode register")
