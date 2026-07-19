# -*- coding: utf-8 -*-
"""Input file 3: Field Validation Reference (.docx). Field ranges/boundaries/enums +
validation order. Drives EP and BVA. No decision logic, no test cases."""
import suite as S
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK=RGBColor(0,0,0); NAVY=RGBColor(0x1F,0x38,0x5C); GREY=RGBColor(0x44,0x44,0x44)
doc=Document()
n=doc.styles["Normal"]; n.font.name="Calibri"; n.font.size=Pt(11); n.font.color.rgb=BLACK
n.paragraph_format.space_after=Pt(6); n.paragraph_format.line_spacing=1.15
rpr=n.element.get_or_add_rPr(); rf=rpr.get_or_add_rFonts(); rf.set(qn('w:eastAsia'),'Calibri')
for s in doc.sections:
    s.top_margin=Inches(0.85); s.bottom_margin=Inches(0.85); s.left_margin=Inches(0.9); s.right_margin=Inches(0.9)

def sr(r,size=11,bold=False,color=BLACK,italic=False):
    r.font.name="Calibri"; r.font.size=Pt(size); r.font.bold=bold; r.font.italic=italic; r.font.color.rgb=color
    rpr=r._element.get_or_add_rPr(); rf=rpr.get_or_add_rFonts(); rf.set(qn('w:eastAsia'),'Calibri')
def para(t,size=11,bold=False,color=BLACK,align=None,after=6,italic=False):
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(after); r=p.add_run(t); sr(r,size,bold,color,italic); return p
def heading(num,text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(f"{num}   {text}"); sr(r,13,True,NAVY)
    pPr=p._p.get_or_add_pPr(); pb=OxmlElement('w:pBdr'); b=OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6'); b.set(qn('w:space'),'2'); b.set(qn('w:color'),'1F385C')
    pb.append(b); pPr.append(pb); return p
def bullet(t,size=11):
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3); r=p.add_run(t); sr(r,size); return p
def setbg(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hexc); tcPr.append(shd)
def table(headers,rows,widths,fs=9):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; r=c.paragraphs[0].add_run(h); sr(r,fs,True,RGBColor(0xFF,0xFF,0xFF)); setbg(c,"1F385C")
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for ci,val in enumerate(row):
            cells[ci].text=""; r=cells[ci].paragraphs[0].add_run(str(val)); sr(r,fs)
            if ri%2==1: setbg(cells[ci],"EEF2F7")
    for row in t.rows:
        for ci,w in enumerate(widths): row.cells[ci].width=Inches(w)
    return t

# cover
p=doc.add_paragraph(); r=p.add_run("PRODUCT REQUIREMENTS"); sr(r,9,True,GREY)
para("LoanPreQual Field Validation Reference",17,True,NAVY,after=2)
para("Companion to LPQ-482 Feature Specification  |  Release 2.4",11,False,GREY,after=2)
para("LPQ-482-VAL",9.5,False,GREY,after=10)
para("This reference defines the validation rule for every field on the applicant record. Validation runs before "
     "any decision rule. A field that fails its rule causes the request to be rejected with the message shown, and "
     "no decision rules run. Fields are validated in the order they appear in Section 2; the first field that fails "
     "is the one reported in the rejection.",11)
para("All numeric ranges below are INCLUSIVE of both stated limits unless the row says otherwise. A value equal to "
     "a stated limit is valid; a value one step beyond it is invalid.",10.5,False,GREY,italic=True)

# validation order
heading("1","Validation Order")
para("Fields are validated in this fixed order. The first failure encountered is the one returned:",11)
for i,f in enumerate(S.FIELD_ORDER,1):
    bullet(f"{i}. {f}")

# numeric fields
heading("2","Field Rules")
para("Numeric fields:",11,bold=True)
rows=[]
labelmap={
 "applicant_age":("Whole years","1 year"),
 "annual_income":("Whole US dollars","1 dollar"),
 "credit_score":("Bureau score, integer","1 point"),
 "loan_amount":("Whole US dollars","1 dollar"),
 "existing_debt_ratio":("Percent, one decimal place","0.1 percent"),
}
for f in ["applicant_age","annual_income","credit_score","loan_amount","existing_debt_ratio"]:
    lo,hi,nom,kind=S.NUM_FIELDS[f]
    unit,stepv=labelmap[f]
    lo_s = f"{lo:.1f}" if kind=="dec1" else f"{lo}"
    hi_s = f"{hi:.1f}" if kind=="dec1" else f"{hi}"
    rows.append([f, unit, f"{lo_s} to {hi_s} inclusive", stepv,
                 f"{f.replace('_',' ')} must be between {lo_s} and {hi_s}"])
table(["Field","Type / format","Valid range","Smallest step","Rejection message"],
      rows,[1.55,1.7,1.7,1.0,2.85])

para("Enumerated fields:",11,bold=True)
erows=[]
for f,members in S.ENUM_FIELDS.items():
    vals=", ".join(str(m) for m in members)
    erows.append([f, vals, f"{f.replace('_',' ')} not a permitted value"])
table(["Field","Permitted values","Rejection message"],erows,[1.9,3.2,3.7])

para("Notes:",11,bold=True)
bullet("existing_debt_ratio is captured to one decimal place. A value with more than one decimal place is "
       "rejected as an invalid format, separately from the range check.")
bullet("annual_income has a lower limit of 0, so a zero income is a valid value for validation purposes. Whether "
       "a zero-income applicant pre-qualifies is a decision-rule question, not a validation question.")
bullet("loan_term_months and employment_status accept only the listed values; anything else is rejected.")

heading("3","What This Reference Does Not Cover")
para("This reference covers field validity only. It does not decide whether a valid applicant is approved, "
     "declined, or referred. Those outcomes are governed entirely by the Decision Rules workbook (LPQ-482-RULES), "
     "which is applied only after every field here passes.",11)

for section in doc.sections:
    fp=section.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=fp.add_run("LPQ-482-VAL Field Validation Reference   |   Page "); sr(r,8,False,GREY)
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); fp._p.append(fld)

cp=doc.core_properties; cp.author="Product Management"; cp.title="LoanPreQual Field Validation Reference"; cp.comments=""
out="inputs/LPQ-482-VAL_Field_Validation_Reference.docx"; doc.save(out); print("saved",out)
