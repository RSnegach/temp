# -*- coding: utf-8 -*-
"""Input file 2: Rs process spec + bin/acceptance definitions (.docx)."""
import wafer as W
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK=RGBColor(0,0,0); NAVY=RGBColor(0x15,0x32,0x4B); GREY=RGBColor(0x44,0x44,0x44)
doc=Document()
n=doc.styles["Normal"]; n.font.name="Calibri"; n.font.size=Pt(11); n.font.color.rgb=BLACK
n.paragraph_format.space_after=Pt(6); n.paragraph_format.line_spacing=1.15
rpr=n.element.get_or_add_rPr(); rf=rpr.get_or_add_rFonts(); rf.set(qn('w:eastAsia'),'Calibri')
for s in doc.sections:
    s.top_margin=Inches(0.9); s.bottom_margin=Inches(0.9); s.left_margin=Inches(1.0); s.right_margin=Inches(1.0)

def sr(r,size=11,bold=False,color=BLACK,italic=False):
    r.font.name="Calibri"; r.font.size=Pt(size); r.font.bold=bold; r.font.italic=italic; r.font.color.rgb=color
    rpr=r._element.get_or_add_rPr(); rf=rpr.get_or_add_rFonts(); rf.set(qn('w:eastAsia'),'Calibri')
def para(t,size=11,bold=False,color=BLACK,align=None,after=6,italic=False):
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(after); r=p.add_run(t); sr(r,size,bold,color,italic); return p
def heading(num,text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(f"{num}   {text}"); sr(r,13,True,NAVY)
    pPr=p._p.get_or_add_pPr(); pb=OxmlElement('w:pBdr'); b=OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6'); b.set(qn('w:space'),'2'); b.set(qn('w:color'),'15324B')
    pb.append(b); pPr.append(pb); return p
def bullet(t,size=11):
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3); r=p.add_run(t); sr(r,size); return p
def setbg(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hexc); tcPr.append(shd)
def table(headers,rows,widths):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; r=c.paragraphs[0].add_run(h); sr(r,9.5,True,RGBColor(0xFF,0xFF,0xFF)); setbg(c,"15324B")
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for ci,val in enumerate(row):
            cells[ci].text=""; r=cells[ci].paragraphs[0].add_run(str(val)); sr(r,9.5)
            if ri%2==1: setbg(cells[ci],"EAF0F4")
    for row in t.rows:
        for ci,w in enumerate(widths): row.cells[ci].width=Inches(w)
    return t

p=doc.add_paragraph(); r=p.add_run("PROCESS ENGINEERING"); sr(r,9,True,GREY)
para("Sheet Resistance Specification and Lot Acceptance",17,True,NAVY,after=2)
para("RTA Anneal, Product Family BX-Logic  |  Post-implant activation",11,False,GREY,after=2)
para("SPEC-RS-118  |  Rev D",9.5,False,GREY,after=10)
para("This specification sets the sheet resistance (Rs) target and limits after the RTA activation anneal, the "
     "within-wafer nonuniformity definition, the site bin thresholds, and the lot acceptance criteria. The site "
     "measurement pattern and the site exclusion and retest rules are defined in the Sampling Plan (SP-RS-045); "
     "this specification does not restate them. Rs values are in ohms per square.",11)
para("CONFIDENTIAL - INTERNAL PROCESS ENGINEERING.",8.5,False,GREY,italic=True)

heading("1","Sheet Resistance Limits")
para(f"The Rs target is {W.RS_TARGET:.0f} ohm/sq. The specification limits are plus or minus {W.RS_TOL_PCT:.0f} "
     "percent of target. A site is within specification when its Rs is greater than or equal to the lower limit "
     "and less than or equal to the upper limit; a site outside either limit is a failing site.",11)
table(["Parameter","Value"],
      [["Target Rs",f"{W.RS_TARGET:.2f} ohm/sq"],
       ["Tolerance",f"+/- {W.RS_TOL_PCT:.0f} % of target"],
       ["Lower spec limit (LSL)",f"{W.RS_LSL:.2f} ohm/sq"],
       ["Upper spec limit (USL)",f"{W.RS_USL:.2f} ohm/sq"]],
      [3.0,3.0])

heading("2","Within-Wafer Nonuniformity")
para("Within-wafer nonuniformity (WIW NU) is computed over the included sites of a wafer, after the site "
     "exclusion rule in the sampling plan has been applied. The definition used for lot acceptance is the "
     "half-range definition:",11)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("WIW NU (%) = (Rs_max - Rs_min) / (2 x Rs_mean) x 100"); sr(r,11,bold=True)
para("where Rs_max, Rs_min, and Rs_mean are taken over the included sites of that wafer. Do not use a standard "
     "deviation based definition for acceptance; the half-range definition above is the controlling one. A "
     "wafer passes the uniformity criterion when its WIW NU is at or below 5.0 percent.",11)

heading("3","Site Bins")
para("Each included site is binned by its Rs relative to target. The warn band flags sites that are within "
     "specification but more than 3 percent from target.",11)
table(["Bin","Condition","Meaning"],
      [["PASS",f"within {W.RS_WARN_LO:.2f} to {W.RS_WARN_HI:.2f} ohm/sq","within 3 % of target"],
       ["WARN",f"in spec but outside the PASS band (>= {W.RS_LSL:.2f} and <= {W.RS_USL:.2f})","3 % to 5 % from target"],
       ["FAIL",f"below {W.RS_LSL:.2f} or above {W.RS_USL:.2f} ohm/sq","out of specification"]],
      [1.1,4.0,1.9])

heading("4","Wafer and Lot Acceptance")
para("Wafer-to-wafer (W2W) variation is computed from the per-wafer mean Rs across the lot, over included sites "
     "only. It is reported as a half-range percentage of the lot mean:",11)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("W2W (%) = (mean_max - mean_min) / (2 x lot_mean) x 100"); sr(r,11,bold=True)
para("Lot yield is the count of included sites within specification across all wafers, divided by the total "
     "count of included sites across all wafers.",11)
bullet("Wafer uniformity criterion: WIW NU at or below 5.0 percent.")
bullet("Wafer-to-wafer criterion: W2W at or below 3.0 percent.")
bullet("Lot yield criterion: at or above 95.0 percent of included sites within specification.")
para("A lot that meets all three criteria is dispositioned CONTINUE. A lot that misses any criterion is placed "
     "on HOLD for engineering review.",11)

for section in doc.sections:
    fp=section.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=fp.add_run("SPEC-RS-118 Rev D   |   Page "); sr(r,8,False,GREY)
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); fp._p.append(fld)

cp=doc.core_properties; cp.author="Process Engineering"; cp.title="Sheet Resistance Specification SPEC-RS-118"; cp.comments=""
out="inputs/SPEC-RS-118_Rs_Specification.docx"; doc.save(out); print("saved",out)
