# -*- coding: utf-8 -*-
"""Input file 3: Sampling Plan (.docx). 49-site pattern, edge-exclusion, retest policy."""
import wafer as W
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK=RGBColor(0,0,0); NAVY=RGBColor(0x15,0x32,0x4B); GREY=RGBColor(0x44,0x44,0x44)
doc=Document()
n=doc.styles["Normal"]; n.font.name="Calibri"; n.font.size=Pt(11); n.font.color.rgb=BLACK
n.paragraph_format.space_after=Pt(6); n.paragraph_format.line_spacing=1.15
rpr=n.element.get_or_add_rPr(); rf=rpr.get_or_add_rFonts(); rf.set(qn('w:eastAsia'),'Calibri')
for s in doc.sections:
    s.top_margin=Inches(0.9); s.bottom_margin=Inches(0.9); s.left_margin=Inches(1.0); s.right_margin=Inches(1.0)

def sr(r,size=11,bold=False,color=BLACK,italic=False):
    r.font.name="Calibri"; r.font.size=Pt(size); r.font.bold=bold; r.font.italic=italic; r.font.color.rgb=color
    rpr=r._element.get_or_add_rPr(); rf=rpr.get_or_add_rFonts(); rf.set(qn('w:eastAsia'),'Calibri')
def para(t,size=11,bold=False,color=BLACK,align=None,after=6,italic=False):
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(after); r=p.add_run(t); sr(r,size,bold,color,italic); return p
def heading(num,text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(f"{num}   {text}"); sr(r,13,True,NAVY)
    pPr=p._p.get_or_add_pPr(); pb=OxmlElement('w:pBdr'); b=OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6'); b.set(qn('w:space'),'2'); b.set(qn('w:color'),'15324B')
    pb.append(b); pPr.append(pb); return p
def bullet(t,size=11):
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3); r=p.add_run(t); sr(r,size); return p
def setbg(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hexc); tcPr.append(shd)
def table(headers,rows,widths):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; r=c.paragraphs[0].add_run(h); sr(r,9.5,True,RGBColor(0xFF,0xFF,0xFF)); setbg(c,"15324B")
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for ci,val in enumerate(row):
            cells[ci].text=""; r=cells[ci].paragraphs[0].add_run(str(val)); sr(r,9.5)
            if ri%2==1: setbg(cells[ci],"EAF0F4")
    for row in t.rows:
        for ci,w in enumerate(widths): row.cells[ci].width=Inches(w)
    return t

p=doc.add_paragraph(); r=p.add_run("PROCESS ENGINEERING"); sr(r,9,True,GREY)
para("Metrology Sampling Plan",17,True,NAVY,after=2)
para("49-Site Sheet Resistance Map  |  150 mm Wafers",11,False,GREY,after=2)
para("SP-RS-045  |  Rev B",9.5,False,GREY,after=10)
para("This plan defines the site sampling pattern for the post-anneal Rs map, the edge exclusion rule, and the "
     "retest handling rule. Spec limits, the nonuniformity definition, and bin thresholds are in SPEC-RS-118. "
     "Site coordinates for each of the 49 sites are provided in the metrology export (Site_Map sheet).",11)
para("CONFIDENTIAL - INTERNAL PROCESS ENGINEERING.",8.5,False,GREY,italic=True)

heading("1","Wafer and Pattern")
para(f"Wafers are 150 mm diameter, so the usable radius is {W.WAFER_RADIUS_MM:.0f} mm from center. The Rs pattern "
     "is a 49-site polar map: one center site and four concentric rings of increasing radius. Sites are numbered "
     "from the center outward, ring by ring. The wafer notch is at the bottom (the negative y direction) and "
     "defines the zero reference for site angles.",11)
table(["Ring","Sites","Nominal radius"],
      [["Center","1","0 mm"],
       ["Ring 1","8","22 mm"],
       ["Ring 2","12","40 mm"],
       ["Ring 3","12","58 mm"],
       ["Ring 4 (outer)","16","73 mm"]],
      [1.8,1.4,2.4])
para("The exact x and y coordinate of every site is listed in the Site_Map sheet of the measurement export. "
     "Use the listed radius for each site rather than assuming the nominal ring radius.",9.5,False,GREY)

heading("2","Edge Exclusion")
para(f"An edge exclusion band of {W.EDGE_EXCLUSION_MM:.0f} mm is applied at the wafer edge. Any site whose radius "
     f"from wafer center is greater than {W.EXCL_RADIUS:.1f} mm falls within the exclusion band and is excluded "
     "from all statistics: within-wafer nonuniformity, wafer mean, wafer-to-wafer, bin counts, and yield. "
     "Excluded sites are still measured and still appear in the export, but they do not enter any calculation.",11)
para("Edge sites carry process artifacts from wafer handling and edge bead and are not representative. Including "
     "them inflates both the mean and the nonuniformity. Confirm each site against its listed radius before "
     "including it.",9.5,False,GREY,italic=True)

heading("3","Retest Handling")
para("A site may be measured more than once when the probe flags a contact or a suspect reading. When a site has "
     "more than one measurement on the same wafer, the measurement with the latest timestamp is the valid one and "
     "supersedes all earlier measurements of that site. Earlier measurements of a retested site are not averaged "
     "in and are not counted; only the latest reading is used.",11)
para("The export does not flag which rows are retests. Identify them by finding sites that appear more than once "
     "for a wafer, and keep the row with the greatest timestamp.",9.5,False,GREY)

heading("4","Order of Operations")
para("Apply the rules in this order before computing any statistic:",11)
bullet("Resolve retests first: for each wafer and site, keep only the latest-timestamp measurement.")
bullet("Apply edge exclusion: drop any site whose radius exceeds the exclusion threshold.")
bullet("Then compute per-wafer statistics, bins, wafer-to-wafer, and lot yield on what remains.")

for section in doc.sections:
    fp=section.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=fp.add_run("SP-RS-045 Rev B   |   Page "); sr(r,8,False,GREY)
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); fp._p.append(fld)

cp=doc.core_properties; cp.author="Process Engineering"; cp.title="Metrology Sampling Plan SP-RS-045"; cp.comments=""
out="inputs/SP-RS-045_Sampling_Plan.docx"; doc.save(out); print("saved",out)
