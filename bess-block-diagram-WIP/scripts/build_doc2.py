# -*- coding: utf-8 -*-
"""Input file 2: Interconnect and Harness Schedule (verbal connection rules, no diagram)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK=RGBColor(0,0,0); NAVY=RGBColor(0x1E,0x2D,0x56); GREY=RGBColor(0x44,0x44,0x44)
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(10.5); normal.font.color.rgb=BLACK
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.15
rpr=normal.element.get_or_add_rPr(); rf=rpr.get_or_add_rFonts(); rf.set(qn('w:eastAsia'),'Arial')
for sec in doc.sections:
    sec.top_margin=Inches(0.9); sec.bottom_margin=Inches(0.9); sec.left_margin=Inches(1.0); sec.right_margin=Inches(1.0)

def sr(r,size=10.5,bold=False,color=BLACK,italic=False):
    r.font.name="Arial"; r.font.size=Pt(size); r.font.bold=bold; r.font.italic=italic; r.font.color.rgb=color
    rpr=r._element.get_or_add_rPr(); rf=rpr.get_or_add_rFonts(); rf.set(qn('w:eastAsia'),'Arial')
def para(t,size=10.5,bold=False,color=BLACK,align=None,after=6,italic=False):
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(after); r=p.add_run(t); sr(r,size,bold,color,italic); return p
def heading(num,text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(f"{num}   {text}"); sr(r,12.5,True,NAVY)
    pPr=p._p.get_or_add_pPr(); pbdr=OxmlElement('w:pBdr'); b=OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6'); b.set(qn('w:space'),'2'); b.set(qn('w:color'),'1E2D56')
    pbdr.append(b); pPr.append(pbdr); return p
def sub(text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(text); sr(r,10.5,True,BLACK); return p
def bullet(t,size=10.5):
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3); r=p.add_run(t); sr(r,size); return p
def setbg(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hexc); tcPr.append(shd)
def table(headers,rows,widths):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,h in enumerate(headers):
        cell=t.rows[0].cells[i]; cell.text=""; p=cell.paragraphs[0]; r=p.add_run(h); sr(r,9,True,RGBColor(0xFF,0xFF,0xFF)); setbg(cell,"1E2D56")
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for ci,val in enumerate(row):
            cells[ci].text=""; p=cells[ci].paragraphs[0]; r=p.add_run(str(val)); sr(r,9)
            if ri%2==1: setbg(cells[ci],"F2F4F8")
    for row in t.rows:
        for ci,w in enumerate(widths): row.cells[ci].width=Inches(w)
    return t

# cover
p=doc.add_paragraph(); r=p.add_run("SYSTEMS ENGINEERING - BESS INTEGRATION"); sr(r,9,True,GREY)
para("Interconnect and Harness Schedule",18,True,NAVY,after=2)
para("Grid-Scale Battery Energy Storage System  |  Standard Product Platform",11,False,GREY,after=2)
para("Document BX-SE-ICD-021  |  Rev 3",9.5,False,GREY,after=10)
para("This schedule defines the four harness families used on the platform and the rules that govern how blocks "
     "interconnect. Block definitions, tags, port complements, and quantities are held in the Block Taxonomy and "
     "Component Register (BX-SE-TAX-014) and are not repeated here. Read the two documents together. This schedule "
     "gives connection logic only; it does not enumerate a specific build.",10.5)
para("Every connection on a diagram belongs to exactly one harness family and is labeled with that family's tag. "
     "The families are drawn in distinct colors on layout diagrams; see the color key in Section 6.",10.5)
para("CONFIDENTIAL - INTERNAL ENGINEERING.",8.5,False,GREY,italic=True)
doc.add_page_break()

# 1 families
heading("1","Harness Families")
table(
    ["Family","Tag","Medium / connector","Carries"],
    [["DC Power","DC-PWR","Cabled DC, lug-terminated","Rack and combiner DC power inside a power block"],
     ["AC Power","AC-PWR","LV/MV AC cable","Power-block AC output and site collection"],
     ["Data / Comms","COMMS","Shielded twisted pair, CAN / Modbus","Rack management data and controller telemetry"],
     ["Aux Control Power","AUX-24V","24 V DC control feed","Regulated control power to active equipment"]],
    [1.35,1.05,2.75,3.35])

# 2 DC-PWR
heading("2","DC Power (DC-PWR)")
para("DC power is wired inside a power block only. It does not cross between power blocks and it does not run to "
     "the site section.",10.5)
bullet("Every battery rack presents its DC output to the DC combiner in its own power block. Each rack makes an "
       "individual run to the combiner; the racks are not looped or chained to one another. This is a star into "
       "the combiner.")
bullet("Each DC combiner makes a single DC run to the power conversion system in the same power block.")
para("So within one power block the DC-PWR count is one run per rack into the combiner, plus one run from the "
     "combiner to the PCS.",9.5,False,GREY)

# 3 AC-PWR
heading("3","AC Power (AC-PWR)")
para("AC power carries each power block's output to the site collection and then to the grid interconnect.",10.5)
bullet("Each power conversion system makes one AC feeder toward the collection.")
bullet("The AC collection bus is a four-position feeder lineup. It can land at most four feeders directly.")
bullet("If the build has four or fewer power blocks, every PCS feeder lands directly on the bus, one per position.")
bullet("If the build has more than four power blocks, the bus cannot take them all directly. In that case the "
       "first three power blocks land directly on the bus, and the remaining power blocks (the fourth and every "
       "one above it) are gathered by the AC combiner panel. Each of those PCS feeders runs to the AC combiner "
       "panel, and the AC combiner panel makes a single feeder onto the fourth bus position.")
bullet("From the collection bus, one AC run goes to revenue metering, and from metering one AC run goes to the MV "
       "step-up transformer.")
para("The AC combiner panel therefore exists only on builds larger than the bus can land directly, and it changes "
     "how the upper power blocks reach the bus. Do not assume AC feeders scale one-for-one with power blocks on "
     "larger builds.",9.5,False,GREY,italic=True)

# 4 COMMS
heading("4","Data and Comms (COMMS)")
para("Comms carries rack data up to the plant controller. Inside a power block the rack management units are "
     "chained; between power blocks and the site there are individual uplinks.",10.5)
bullet("Within a power block the rack BMS units are connected in a daisy chain: the first BMS connects to the "
       "second, the second to the third, and so on down the block. They are not each wired back to the controller "
       "and they are not wired to the combiner.")
bullet("The head of the chain (the first BMS in the power block) makes one comms run to the power conversion "
       "system in that block. The chain reports to the site only through the PCS; the tail BMS has no onward run.")
bullet("Each power conversion system makes one comms run to the site controller. This is an individual uplink per "
       "power block, not a shared bus.")
bullet("Revenue metering makes its own comms run to the site controller.")
para("So a power block's comms count is the chain links between its BMS units, plus one head-to-PCS run, plus one "
     "PCS-to-controller uplink. Metering adds one site-level run.",9.5,False,GREY)

# 5 AUX-24V
heading("5","Aux Control Power (AUX-24V)")
para("The auxiliary transformer distributes 24 V control power in a star to the equipment that needs it.",10.5)
bullet("The auxiliary transformer feeds each power conversion system and each DC combiner. That is two aux runs "
       "per power block.")
bullet("The auxiliary transformer also feeds the site controller. That is one site-level aux run.")
bullet("Battery racks and rack BMS units are not aux-fed; they take their control power from their own DC. The AC "
       "collection bus, revenue metering, the MV step-up transformer, and the AC combiner panel are not aux-fed "
       "either.")
para("So aux runs are two per power block plus one at the site. Do not add an aux run for racks or BMS.",9.5,False,GREY)

# 6 color key + drawing conventions
heading("6","Diagram Conventions")
para("On a layout or block diagram, draw each block as a labeled box and each connection as a line tagged with "
     "its harness family. Use the following colors so families read at a glance:",10.5)
table(["Harness family","Line color"],
      [["DC-PWR","Red"],["AC-PWR","Black"],["COMMS","Blue"],["AUX-24V","Green"]],
      [3.0,2.0])
para("Every line on the diagram must carry its family tag. A connection with no tag, or a line whose color does "
     "not match its tag, is a drafting error.",9.5,False,GREY)

for section in doc.sections:
    fp=section.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=fp.add_run("BX-SE-ICD-021 Rev 3   |   Page "); sr(r,8,False,GREY)
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); fp._p.append(fld)

cp=doc.core_properties; cp.author="Systems Engineering"; cp.title="BESS Interconnect and Harness Schedule"; cp.comments=""
out="inputs/BX-SE-ICD-021_Interconnect_Harness_Schedule.docx"; doc.save(out); print("saved",out)
