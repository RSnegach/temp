# -*- coding: utf-8 -*-
"""Input file 1: BESS block taxonomy (verbal description of every block type)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK=RGBColor(0,0,0); NAVY=RGBColor(0x1E,0x2D,0x56); GREY=RGBColor(0x44,0x44,0x44)
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(10.5); normal.font.color.rgb=BLACK
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.15
rpr=normal.element.get_or_add_rPr(); rf=rpr.get_or_add_rFonts(); rf.set(qn('w:eastAsia'),'Arial')
for sec in doc.sections:
    sec.top_margin=Inches(0.9); sec.bottom_margin=Inches(0.9); sec.left_margin=Inches(1.0); sec.right_margin=Inches(1.0)

def sr(r,size=10.5,bold=False,color=BLACK,italic=False):
    r.font.name="Arial"; r.font.size=Pt(size); r.font.bold=bold; r.font.italic=italic; r.font.color.rgb=color
    rpr=r._element.get_or_add_rPr(); rf=rpr.get_or_add_rFonts(); rf.set(qn('w:eastAsia'),'Arial')
def para(t,size=10.5,bold=False,color=BLACK,align=None,after=6,italic=False):
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(after); r=p.add_run(t); sr(r,size,bold,color,italic); return p
def heading(num,text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(f"{num}   {text}"); sr(r,12.5,True,NAVY)
    pPr=p._p.get_or_add_pPr(); pbdr=OxmlElement('w:pBdr'); b=OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6'); b.set(qn('w:space'),'2'); b.set(qn('w:color'),'1E2D56')
    pbdr.append(b); pPr.append(pbdr); return p
def bullet(t,size=10.5):
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3); r=p.add_run(t); sr(r,size); return p
def setbg(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hexc); tcPr.append(shd)
def table(headers,rows,widths):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,h in enumerate(headers):
        cell=t.rows[0].cells[i]; cell.text=""; p=cell.paragraphs[0]; r=p.add_run(h); sr(r,9,True,RGBColor(0xFF,0xFF,0xFF)); setbg(cell,"1E2D56")
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for ci,val in enumerate(row):
            cells[ci].text=""; p=cells[ci].paragraphs[0]; r=p.add_run(str(val)); sr(r,9)
            if ri%2==1: setbg(cells[ci],"F2F4F8")
    for row in t.rows:
        for ci,w in enumerate(widths): row.cells[ci].width=Inches(w)
    return t

# cover
p=doc.add_paragraph(); r=p.add_run("SYSTEMS ENGINEERING - BESS INTEGRATION"); sr(r,9,True,GREY)
para("Block Taxonomy and Component Register",18,True,NAVY,after=2)
para("Grid-Scale Battery Energy Storage System  |  Standard Product Platform",11,False,GREY,after=2)
para("Document BX-SE-TAX-014  |  Rev 3",9.5,False,GREY,after=10)
para("This register defines every block type that appears on a site single-line block diagram for the "
     "standard containerized BESS platform. It gives the function, tag convention, port complement, and the "
     "quantity rule for each block type. It does not define how blocks interconnect; interconnect and harness "
     "rules are held separately in the Interconnect and Harness Schedule (BX-SE-ICD-021).",10.5)
para("CONFIDENTIAL - INTERNAL ENGINEERING. Platform reference for proposal and layout work.",8.5,False,GREY,italic=True)
doc.add_page_break()

# 1 architecture overview
heading("1","Platform Architecture")
para("The platform is organized into a shared site section and a repeating power block. The site section is "
     "common to every project and is built once regardless of how many power blocks the project carries. The "
     "power block is the unit of scale: a project is quoted as a number of power blocks, and every power block "
     "is identical in its internal makeup.",10.5)
para("Each power block carries one power conversion system, one DC combiner, and a fixed complement of four "
     "battery racks. Every battery rack has its own rack battery management unit. The site section carries the "
     "site controller, the auxiliary transformer, the medium-voltage step-up transformer, revenue metering, and "
     "the AC collection bus. One further site block, the AC combiner panel, is fitted only on larger builds; see "
     "its entry below for the fitment rule.",10.5)
para("Tag convention: site blocks carry a fixed tag. Power-block equipment is numbered by power block. Where a "
     "block repeats inside a power block (racks and their BMS), the tag carries the power block number and then "
     "the position within the block. So the third rack in power block two is RK-2-3, and its management unit is "
     "BMS-2-3. Power conversion systems and DC combiners are numbered by power block only, zero-padded to two "
     "digits, for example PCS-02 and DCC-02.",10.5)

# 2 site blocks
heading("2","Site Blocks (built once per site)")
table(
    ["Block type","Tag","Function","Ports","Qty rule"],
    [["Site Controller","SC-01","Plant controller and SCADA gateway. Aggregates telemetry, dispatches setpoints.",
      "COMMS (uplinks in), AUX-24V (in)","One per site"],
     ["Aux Transformer","AUX-XFMR","Derives regulated 24 V control power for the site from the collection bus.",
      "AUX-24V (out, multiple)","One per site"],
     ["MV Step-up Transformer","MV-XFMR","Steps the collection voltage up to the medium-voltage point of interconnect.",
      "AC-PWR (in)","One per site"],
     ["Revenue Metering","MTR-01","Revenue-grade metering at the point of common coupling.",
      "AC-PWR (in, out), COMMS (out)","One per site"],
     ["AC Collection Bus","AC-BUS","Low-voltage AC lineup that collects the power-block feeders. Four feeder positions.",
      "AC-PWR (in x4, out)","One per site"],
     ["AC Combiner Panel","AC-CMB","Aggregates surplus power-block feeders into a single bus feeder when the site "
      "exceeds the bus feeder count.","AC-PWR (in multiple, out one)","Fitted only when the site has more power "
      "blocks than the AC collection bus has feeder positions; otherwise omitted"]],
    [1.35,1.0,3.1,1.6,1.35])

# 3 power block
heading("3","Power Block Equipment (repeats per power block)")
para("A power block is the repeating unit. Every power block is identical: one PCS, one DC combiner, four battery "
     "racks, and one rack BMS per rack. There are no partial power blocks.",10.5)
table(
    ["Block type","Tag pattern","Function","Ports","Qty per power block"],
    [["Power Conversion System","PCS-nn","Bidirectional inverter for the power block. AC side to collection, DC "
      "side to the combiner.","DC-PWR (in), AC-PWR (out), COMMS (in from block, out to site), AUX-24V (in)","One"],
     ["DC Combiner","DCC-nn","Combines four incoming rack DC feeds onto a single DC output.",
      "DC-PWR (in x4, out), AUX-24V (in)","One"],
     ["Battery Rack","RK-n-k","Energy storage rack. Presents a single DC power output.",
      "DC-PWR (out)","Four"],
     ["Rack BMS","BMS-n-k","Rack management and protection unit. One per rack. Reports on the block data segment.",
      "COMMS (in, out)","Four (one per rack)"]],
    [1.6,1.05,2.9,1.9,1.3])
para("Note on the rack BMS: the BMS is drawn as its own block on the diagram, adjacent to the rack it manages. It "
     "is not folded into the rack symbol.",9.5,False,GREY)

# 4 quantities
heading("4","Quantity Summary")
para("The site section is fixed. The power-block equipment scales with the number of power blocks in the build. "
     "The AC combiner panel is the one site block whose presence depends on the build size, per its fitment rule "
     "in Section 2. Rack and BMS quantities follow from four racks per power block.",10.5)
bullet("Site controller, auxiliary transformer, MV step-up transformer, revenue metering, AC collection bus: one each, every build.")
bullet("Per power block: one PCS, one DC combiner, four battery racks, four rack BMS.")
bullet("AC combiner panel: present only when the number of power blocks exceeds the AC collection bus feeder count.")

# footer page numbers
for section in doc.sections:
    fp=section.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=fp.add_run("BX-SE-TAX-014 Rev 3   |   Page "); sr(r,8,False,GREY)
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); fp._p.append(fld)

cp=doc.core_properties; cp.author="Systems Engineering"; cp.title="BESS Block Taxonomy and Component Register"; cp.comments=""
out="inputs/BX-SE-TAX-014_Block_Taxonomy.docx"; doc.save(out); print("saved",out)
