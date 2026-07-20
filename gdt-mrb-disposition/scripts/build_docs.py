# -*- coding: utf-8 -*-
"""
Build the GD&T interpretation and disposition-procedure document (input file).
This is the professionally formatted reference that states HOW to disposition:
bonus tolerance, datum shift, virtual condition, the order of checks, and the
disposition rule. It does NOT contain the answers; the engineer applies it to
the drawing callouts and the CMM report.

Formatting per current Geranium rules: all text black, neutral light-gray table
headers, no colored accent, no italic subtitle rows, no em dashes.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0, 0, 0)
HDRFILL = "D9D9D9"

def newdoc():
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(11)
    n.font.color.rgb = BLACK
    n.paragraph_format.space_after = Pt(6); n.paragraph_format.line_spacing = 1.15
    rpr = n.element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts(); rf.set(qn('w:eastAsia'), 'Calibri')
    for s in doc.sections:
        s.top_margin = Inches(0.9); s.bottom_margin = Inches(0.9)
        s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)
    return doc

def sr(r, size=11, bold=False, italic=False):
    r.font.name = "Calibri"; r.font.size = Pt(size); r.font.bold = bold
    r.font.italic = italic; r.font.color.rgb = BLACK
    rpr = r._element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts(); rf.set(qn('w:eastAsia'), 'Calibri')

def para(doc, t, size=11, bold=False, after=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(t); sr(r, size, bold); return p

def heading(doc, num, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{num}   {text}"); sr(r, 13, True)
    pPr = p._p.get_or_add_pPr(); pb = OxmlElement('w:pBdr'); b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6'); b.set(qn('w:space'), '2'); b.set(qn('w:color'), '000000')
    pb.append(b); pPr.append(pb); return p

def bullet(doc, t, size=11):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(t); sr(r, size); return p

def setbg(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexc)
    tcPr.append(shd)

def table(doc, headers, rows, widths, fs=9.5):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(h); sr(r, fs, True); setbg(c, HDRFILL)
    for row in rows:
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""; r = cells[ci].paragraphs[0].add_run(str(val)); sr(r, fs)
    for row in t.rows:
        for ci, w in enumerate(widths):
            row.cells[ci].width = Inches(w)
    return t

def main(out="GDT-DISP_Interpretation_and_Disposition_Procedure.docx"):
    doc = newdoc()
    para(doc, "QUALITY ENGINEERING, DIMENSIONAL", 9, True, after=2)
    para(doc, "GD&T Interpretation and Inspection Disposition Procedure", 17, True, after=2)
    para(doc, "Position tolerancing with material condition modifiers  |  ASME Y14.5 basis", 11, False, after=2)
    para(doc, "DOC GDT-DISP Rev B", 9.5, False, after=10)
    para(doc, "This procedure defines how to disposition a machined feature from its drawing "
              "callout and its CMM measurement. It gives the interpretation rules and the order of "
              "operations. It does not contain part-specific results. Apply it to the feature control "
              "frames on the drawing and the measured values in the CMM report to reach a disposition "
              "for each feature.")

    heading(doc, "1", "Order of Operations")
    para(doc, "Disposition each feature in this order. Do not skip ahead: a feature that fails an "
              "earlier check is dispositioned there and later checks do not apply.")
    bullet(doc, "Step 1. Size. Confirm the measured local size is within the size limits. A feature "
                "outside its size limits is rejected on size and its position is not evaluated.")
    bullet(doc, "Step 2. Bonus tolerance. If the position callout carries a maximum material condition "
                "modifier, compute the bonus from the feature's departure from MMC.")
    bullet(doc, "Step 3. Datum shift. If a datum feature of size is referenced at MMC in the frame, add "
                "that datum's departure from its own MMC as additional allowance.")
    bullet(doc, "Step 4. Total allowed position tolerance = stated tolerance + bonus + datum shift.")
    bullet(doc, "Step 5. Compare. Convert the measured center offset to a diameter and compare to the "
                "total allowed. Assign the disposition.")

    heading(doc, "2", "Feature Control Frame Reading")
    para(doc, "A position feature control frame reads left to right: the position symbol, then the "
              "tolerance value (a diameter zone when preceded by the diameter sign), then any material "
              "condition modifier on the tolerance, then the datum references in precedence order. A "
              "modifier may also follow a datum letter, which applies the material condition to that "
              "datum.")
    table(doc,
          ["Symbol or mark", "Meaning"],
          [["Circled plus", "Position: the axis or center must lie within a tolerance zone about true position"],
           ["Diameter sign before the value", "The position tolerance zone is a diameter (cylindrical), not a width"],
           ["Circled M on the tolerance", "Maximum material condition (MMC): bonus tolerance is available as the feature departs from MMC"],
           ["Circled M on a datum letter", "That datum feature of size is taken at MMC: datum shift is available"],
           ["No modifier on the tolerance", "Regardless of feature size (RFS): no bonus tolerance, no matter how the size departs"],
           ["Arc on a baseline", "Profile of a surface: the surface must lie within the stated zone; not a feature of size"],
           ["Parallelogram", "Flatness: the surface must lie within the stated zone; not a feature of size"]],
          [2.2, 4.6])

    heading(doc, "3", "Maximum Material Condition and Bonus Tolerance")
    para(doc, "MMC is the size limit at which a feature contains the most material. For a hole (internal "
              "feature) MMC is the smallest allowed size. For a pin or boss (external feature) MMC is the "
              "largest allowed size. When a position tolerance is specified at MMC, the feature earns bonus "
              "tolerance equal to its departure from MMC.")
    bullet(doc, "Internal feature (hole): bonus = measured local size minus MMC size. The hole earns bonus "
                "as it grows larger than MMC.")
    bullet(doc, "External feature (pin): bonus = MMC size minus measured local size. The pin earns bonus as "
                "it grows smaller than MMC.")
    bullet(doc, "At MMC exactly, the departure is zero, so the bonus is zero.")
    bullet(doc, "RFS (no modifier): bonus is always zero, regardless of how far the size departs from any limit.")

    heading(doc, "4", "Datum Shift")
    para(doc, "When a datum feature of size is referenced at MMC in the feature control frame, the part is "
              "allowed to shift relative to that datum by the datum feature's own departure from its MMC. "
              "This datum shift adds to the allowed position tolerance for every feature controlled to that "
              "datum at MMC. A datum referenced without a modifier (RFS) provides no shift. A planar datum "
              "(a flat face, not a feature of size) never provides shift.")
    bullet(doc, "Datum shift (internal datum feature) = measured datum size minus datum MMC.")
    bullet(doc, "Datum shift (external datum feature) = datum MMC minus measured datum size.")
    bullet(doc, "Add the datum shift only when the frame shows the modifier on that datum letter.")

    heading(doc, "5", "Actual Position Deviation")
    para(doc, "The CMM reports the feature center offset from true position as two components, one along "
              "each datum axis. The positional deviation as a diameter is:")
    para(doc, "     positional deviation (diameter) = 2 times the square root of ( devX squared + devY squared )",
         size=11, bold=True)
    para(doc, "Compare this diameter directly to the total allowed position tolerance, which is also a "
              "diameter. Do not compare a radius to a diameter.")

    heading(doc, "6", "Virtual Condition (reference)")
    para(doc, "Virtual condition is the constant mating boundary the feature must respect. It is reported "
              "for reference; the accept or reject decision is made on the position deviation versus the "
              "total allowed tolerance, not on virtual condition.")
    bullet(doc, "Internal feature: virtual condition = MMC size minus the stated geometric tolerance.")
    bullet(doc, "External feature: virtual condition = MMC size plus the stated geometric tolerance.")

    heading(doc, "7", "Disposition Rule")
    para(doc, "Assign exactly one disposition to each feature:")
    table(doc,
          ["Condition", "Disposition"],
          [["Measured size outside the size limits", "REJECT (size)"],
           ["Size in limits and position deviation is within the total allowed tolerance", "ACCEPT"],
           ["Size in limits and position deviation exceeds the total allowed tolerance", "MRB (material review)"],
           ["Profile or flatness deviation within the stated tolerance", "ACCEPT"],
           ["Profile or flatness deviation exceeds the stated tolerance", "MRB (material review)"]],
          [4.4, 2.4])
    para(doc, "Profile and flatness are not features of size. They carry no bonus and no datum shift; "
              "compare the measured surface deviation directly to the stated tolerance.")

    heading(doc, "8", "Deliverable")
    para(doc, "Return one row per feature showing the extracted callout, the stated tolerance, the bonus, "
              "the datum shift, the total allowed tolerance, the actual deviation, the size check, and the "
              "final disposition, plus a count of features in each disposition category. Compute the derived "
              "values with live spreadsheet formulas that reference the measured inputs, so the workbook "
              "recalculates if a measurement is corrected.")

    doc.save(out)
    print("wrote", out)

if __name__ == "__main__":
    main()
