# -*- coding: utf-8 -*-
"""Input file 1: LoanPreQual feature specification (.docx). Narrative + acceptance
criteria, no test cases. Field-authentic product/BA spec voice."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK=RGBColor(0,0,0); NAVY=RGBColor(0x1F,0x38,0x5C); GREY=RGBColor(0x44,0x44,0x44)
doc=Document()
n=doc.styles["Normal"]; n.font.name="Calibri"; n.font.size=Pt(11); n.font.color.rgb=BLACK
n.paragraph_format.space_after=Pt(6); n.paragraph_format.line_spacing=1.15
rpr=n.element.get_or_add_rPr(); rf=rpr.get_or_add_rFonts(); rf.set(qn('w:eastAsia'),'Calibri')
for s in doc.sections:
    s.top_margin=Inches(0.9); s.bottom_margin=Inches(0.9); s.left_margin=Inches(1.0); s.right_margin=Inches(1.0)

def sr(r,size=11,bold=False,color=BLACK,italic=False):
    r.font.name="Calibri"; r.font.size=Pt(size); r.font.bold=bold; r.font.italic=italic; r.font.color.rgb=color
    rpr=r._element.get_or_add_rPr(); rf=rpr.get_or_add_rFonts(); rf.set(qn('w:eastAsia'),'Calibri')
def para(t,size=11,bold=False,color=BLACK,align=None,after=6,italic=False):
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(after); r=p.add_run(t); sr(r,size,bold,color,italic); return p
def heading(num,text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(f"{num}   {text}"); sr(r,13,True,NAVY)
    pPr=p._p.get_or_add_pPr(); pb=OxmlElement('w:pBdr'); b=OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6'); b.set(qn('w:space'),'2'); b.set(qn('w:color'),'1F385C')
    pb.append(b); pPr.append(pb); return p
def sub(text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(text); sr(r,11,True,BLACK); return p
def bullet(t,size=11):
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3); r=p.add_run(t); sr(r,size); return p
def num_item(t,size=11):
    p=doc.add_paragraph(style="List Number"); p.paragraph_format.space_after=Pt(3); r=p.add_run(t); sr(r,size); return p
def setbg(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hexc); tcPr.append(shd)
def table(headers,rows,widths):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; r=c.paragraphs[0].add_run(h); sr(r,9.5,True,RGBColor(0xFF,0xFF,0xFF)); setbg(c,"1F385C")
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for ci,val in enumerate(row):
            cells[ci].text=""; r=cells[ci].paragraphs[0].add_run(str(val)); sr(r,9.5)
            if ri%2==1: setbg(cells[ci],"EEF2F7")
    for row in t.rows:
        for ci,w in enumerate(widths): row.cells[ci].width=Inches(w)
    return t

# cover
p=doc.add_paragraph(); r=p.add_run("PRODUCT REQUIREMENTS"); sr(r,9,True,GREY)
para("LoanPreQual Decision Service",18,True,NAVY,after=2)
para("Feature Specification  |  Consumer Lending Platform",11,False,GREY,after=2)
para("JIRA LPQ-482  |  Release 2.4  |  Status: Approved for build",9.5,False,GREY,after=10)
para("This specification defines the behaviour of the pre-qualification decision service. The service takes a "
     "single applicant record and returns one decision. It is a synchronous call from the online application "
     "front end and from the branch origination tool. This document is the source of truth for behaviour. Field "
     "validation limits are held in the Field Validation Reference (LPQ-482-VAL), and the decision rule set and "
     "its evaluation order are held in the Decision Rules workbook (LPQ-482-RULES). Read all three together.",11)
para("This document does not contain test cases. QA will derive the test suite from this specification and the "
     "two companion references.",10.5,False,GREY,italic=True)
doc.add_page_break()

# 1 scope
heading("1","Purpose and Scope")
para("LoanPreQual gives an applicant an immediate, non-binding pre-qualification decision for an unsecured "
     "personal loan. It does not pull a hard credit file, disburse funds, or create an account. It reads the "
     "applicant record captured on the intake form, validates it, evaluates the lending rules, and returns one "
     "of a fixed set of decision outcomes with a short reason.",11)
para("In scope: field validation of the applicant record, evaluation of the decision rule set, and the returned "
     "decision. Out of scope: identity verification, fraud scoring, document upload, and downstream funding. "
     "Those are separate services and are not exercised here.",11)

# 2 request
heading("2","Applicant Record (Request)")
para("The service accepts one applicant record per call. Every field is mandatory. The fields, their business "
     "meaning, and their permitted values are specified in the Field Validation Reference; the summary below is "
     "for context only and is not the authority on limits.",11)
table(["Field","Meaning"],
      [["applicant_age","Age of the primary applicant in whole years."],
       ["annual_income","Gross annual income in whole US dollars."],
       ["credit_score","Soft-pull bureau score."],
       ["loan_amount","Requested principal in whole US dollars."],
       ["existing_debt_ratio","Current debt-to-income ratio as a percentage, to one decimal place."],
       ["employment_status","Primary applicant employment category."],
       ["loan_term_months","Requested repayment term."]],
      [1.9,4.6])

# 3 processing order
heading("3","Processing Sequence")
para("The service processes a request in two stages, in this order:",11)
num_item("Validation. Every field is checked against the Field Validation Reference. If one or more fields fail "
         "validation, the service returns a rejection that names the first failing field and its message, and no "
         "decision rules are evaluated. Fields are validated in the order listed in the reference, and the first "
         "failure is the one reported.")
num_item("Decision. If and only if every field is valid, the applicant record is passed to the decision rule "
         "engine. The engine evaluates the rules in the fixed order given in the Decision Rules workbook and "
         "returns the outcome of the first rule that matches. Exactly one decision is returned per valid record.")
para("A field that is out of range is a validation rejection, not a decline. A decline is a business decision "
     "made on a fully valid record. These are different outcomes and must not be conflated.",10.5,False,GREY,italic=True)

# 4 outcomes
heading("4","Decision Outcomes")
para("Every call returns exactly one of the outcomes below.",11)
table(["Outcome","Meaning"],
      [["REJECT","One or more fields failed validation. The response names the first failing field."],
       ["DECLINE","The record is valid but a business rule denies pre-qualification."],
       ["REFER","The record is valid but falls in the manual underwriting band; no automated approve or decline."],
       ["APPROVE_STANDARD","Pre-qualified at the standard tier."],
       ["APPROVE_PRIME","Pre-qualified at the prime tier, which carries the best rate card."]],
      [1.9,4.6])

# 5 acceptance criteria (behavioural, not the rule thresholds themselves)
heading("5","Acceptance Criteria")
para("The following acceptance criteria are written in business terms. The exact numeric thresholds live in the "
     "Decision Rules workbook; do not assume a threshold value from this section.",11)
bullet("AC-1: A record with any field outside its permitted range returns REJECT and names the first failing "
       "field, and no decision rule is evaluated.")
bullet("AC-2: An applicant with no qualifying income source is declined regardless of any other field.")
bullet("AC-3: An applicant whose bureau score is below the lending floor is declined, unless an earlier rule "
       "already applies.")
bullet("AC-4: An applicant whose existing debt-to-income exceeds the program maximum is declined, unless an "
       "earlier rule already applies.")
bullet("AC-5: A requested principal that is large relative to income is declined, unless an earlier rule "
       "already applies. The relationship is a multiple of annual income; the multiple is in the rules workbook.")
bullet("AC-6: A strong-credit, low-debt applicant is approved at the prime tier.")
bullet("AC-7: An applicant who clears the lending floor but is not prime is approved at the standard tier.")
bullet("AC-8: An applicant who sits above the lending floor but below the standard-approval band is referred "
       "to manual underwriting rather than auto-declined.")
para("Rule precedence matters. Where more than one rule could apply, the earlier rule in the evaluation order "
     "wins, and the later condition is never reached. Test design must respect that a record can satisfy several "
     "conditions at once and that only the first is returned.",10.5,False,GREY,italic=True)

# 6 notes
heading("6","Notes for Test Design")
para("Boundary behaviour is significant for this feature. Where the rules workbook states a threshold, the "
     "behaviour at the threshold value, one unit below, and one unit above is defined by whether the comparison "
     "is inclusive or strict; the workbook states the operator for each threshold. The validation reference "
     "likewise states whether each range limit is inclusive. Test coverage is expected to exercise each "
     "validation range limit and each decision threshold, not just nominal values.",11)

for section in doc.sections:
    fp=section.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=fp.add_run("LPQ-482 Feature Specification, Release 2.4   |   Page "); sr(r,8,False,GREY)
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); fp._p.append(fld)

cp=doc.core_properties; cp.author="Product Management"; cp.title="LoanPreQual Decision Service Feature Specification"; cp.comments=""
out="inputs/LPQ-482_Feature_Specification.docx"; doc.save(out); print("saved",out)
